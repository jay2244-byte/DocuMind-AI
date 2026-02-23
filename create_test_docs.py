import docx
from fpdf import FPDF
import os

sample_text = """
The Solar System is the gravitationally bound system of the Sun and the objects that orbit it. 
It formed 4.6 billion years ago from the gravitational collapse of a giant interstellar molecular cloud. 
The vast majority of the system's mass is in the Sun, with most of the remaining mass contained in Jupiter.
"""

def create_docx():
    doc = docx.Document()
    doc.add_heading('Solar System Overview', 0)
    doc.add_paragraph(sample_text)
    doc.save('sample.docx')
    print("Created sample.docx")

def create_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Solar System Overview", ln=1, align='C')
    pdf.multi_cell(0, 10, txt=sample_text)
    pdf.output("sample.pdf")
    print("Created sample.pdf")

def create_txt():
    with open('sample.txt', 'w', encoding='utf-8') as f:
        f.write(sample_text)
    print("Created sample.txt")

if __name__ == "__main__":
    create_docx()
    # FPDF is not in requirements.txt, so I will only do DOCX and TXT here
    create_txt()
